# -*- coding: utf-8 -*-
"""
report_generator.py  —  PDF + DOCX Stock Analysis Report Generator
====================================================================
Generates professional downloadable reports for any stock.

Dependencies (add to requirements.txt):
    reportlab
    python-docx

Routes added to app.py:
    GET /api/report/<symbol>?format=pdf   → download PDF
    GET /api/report/<symbol>?format=docx  → download DOCX
    GET /api/report/<symbol>              → download PDF (default)

Usage in app.py:
    from report_generator import register_report_routes
    register_report_routes(app, build_stock_context)
"""

import io, os, time, datetime
from flask import send_file, jsonify, request

# ── Graceful import of reportlab ─────────────────────────────────────────────
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                     TableStyle, HRFlowable, PageBreak, KeepTogether)
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

# ── Graceful import of python-docx ───────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


# ── Colours ───────────────────────────────────────────────────────────────────
BRAND     = colors.HexColor("#38bdf8")
TEAL      = colors.HexColor("#2dd4bf")
GREEN     = colors.HexColor("#10b981")
RED       = colors.HexColor("#f87171")
AMBER     = colors.HexColor("#fbbf24")
DARK_BG   = colors.HexColor("#0c1628")
MID_BG    = colors.HexColor("#0f1e35")
LIGHT_BG  = colors.HexColor("#1a2f50")
TEXT      = colors.HexColor("#e8f0ff")
MUTED     = colors.HexColor("#5a7090")
WHITE     = colors.white
BLACK     = colors.black

def _money(v):
    if v is None or v == "" or v == "N/A": return "—"
    try: return "₹{:,.2f}".format(float(v))
    except: return str(v)

def _pct(v, suffix="%"):
    if v is None or v == "": return "—"
    try:
        fv = float(v)
        return f"+{fv:.2f}{suffix}" if fv >= 0 else f"{fv:.2f}{suffix}"
    except: return str(v)

def _fmt(v, suffix=""):
    if v is None or str(v).strip() in ("", "N/A", "0"): return "—"
    return f"{v}{suffix}"

def _sig_color(sig):
    if not sig: return AMBER
    s = sig.upper()
    if "STRONG BUY" in s or "BUY" in s:   return GREEN
    if "STRONG SELL" in s or "SELL" in s: return RED
    return AMBER

# ═════════════════════════════════════════════════════════════════════════════
# PDF GENERATOR
# ═════════════════════════════════════════════════════════════════════════════

def generate_pdf(context: dict) -> io.BytesIO:
    """Build a multi-page PDF report from the stock context dict."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=1.8*cm, leftMargin=1.8*cm,
        topMargin=2.0*cm,   bottomMargin=2.0*cm,
        title=f"Stock Report — {context.get('symbol','')}"
    )

    styles  = getSampleStyleSheet()
    story   = []

    # ── Custom styles ─────────────────────────────────────────────────────────
    def S(name, parent="Normal", **kw):
        return ParagraphStyle(name, parent=styles[parent], **kw)

    TITLE      = S("T",  fontSize=24, textColor=WHITE, fontName="Helvetica-Bold",
                   spaceAfter=4, leading=28)
    SUBTITLE   = S("St", fontSize=13, textColor=BRAND,  fontName="Helvetica",
                   spaceAfter=8)
    H1         = S("H1", fontSize=13, textColor=BRAND,  fontName="Helvetica-Bold",
                   spaceBefore=14, spaceAfter=6)
    H2         = S("H2", fontSize=11, textColor=TEXT,   fontName="Helvetica-Bold",
                   spaceBefore=8,  spaceAfter=4)
    NORMAL     = S("N",  fontSize=9,  textColor=TEXT,   fontName="Helvetica",
                   leading=14, spaceAfter=3)
    MUTED_S    = S("M",  fontSize=8,  textColor=MUTED,  fontName="Helvetica",
                   leading=12)
    VERDICT    = S("V",  fontSize=18, fontName="Helvetica-Bold", spaceAfter=6)
    DISCLAIMER = S("D",  fontSize=7,  textColor=MUTED,  fontName="Helvetica",
                   leading=10)

    # ── HEADER BANNER ─────────────────────────────────────────────────────────
    sym  = context.get("symbol",   "N/A")
    name = context.get("name",     sym)
    exch = context.get("exchange", "NSE")
    ts   = datetime.datetime.now().strftime("%d %b %Y, %I:%M %p")

    header_data = [[
        Paragraph(f"<b>{name}</b>", S("HN", fontSize=20, textColor=WHITE, fontName="Helvetica-Bold")),
        Paragraph(f"<b>{sym}</b><br/><font size='9' color='#5a7090'>{exch} · Generated {ts}</font>",
                  S("HR", fontSize=12, textColor=BRAND, fontName="Helvetica-Bold", alignment=TA_RIGHT))
    ]]
    ht = Table(header_data, colWidths=["60%", "40%"])
    ht.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,-1), DARK_BG),
        ("ROWPADDING",   (0,0), (-1,-1), 14),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
        ("LINEBELOW",    (0,0), (-1,-1), 1.5, BRAND),
        ("ROUNDEDCORNERS", [6]),
    ]))
    story.append(ht)
    story.append(Spacer(1, 10))

    # ── PRICE + SIGNAL SUMMARY ─────────────────────────────────────────────────
    price = context.get("price", {}) or {}
    tech  = context.get("technical", {}) or {}
    fund  = context.get("fundamental", {}) or {}
    verd  = context.get("verdict", {}) or {}
    news  = context.get("news", [])

    sig      = tech.get("signal", verd.get("verdict", "HOLD"))
    sig_col  = _sig_color(sig)
    p        = price.get("price")
    pct_chg  = price.get("pct") or price.get("change_pct")
    pct_col  = GREEN if (pct_chg or 0) >= 0 else RED

    kpi_data = [
        ["PRICE",         "DAY CHANGE",    "SIGNAL",     "BULL SCORE"],
        [
            Paragraph(f"<b>{_money(p)}</b>", S("KV", fontSize=16, textColor=WHITE, fontName="Helvetica-Bold", alignment=TA_CENTER)),
            Paragraph(f"<b>{_pct(pct_chg)}</b>", S("KC", fontSize=14, fontName="Helvetica-Bold",
                       textColor=pct_col, alignment=TA_CENTER)),
            Paragraph(f"<b>{sig}</b>", S("KS", fontSize=14, fontName="Helvetica-Bold",
                       textColor=sig_col, alignment=TA_CENTER)),
            Paragraph(f"<b>{_fmt(tech.get('bull_score'),'%')}</b>",
                      S("KB", fontSize=14, textColor=BRAND, fontName="Helvetica-Bold", alignment=TA_CENTER)),
        ]
    ]
    kpi_t = Table(kpi_data, colWidths=["25%","25%","25%","25%"])
    kpi_t.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,0), LIGHT_BG),
        ("BACKGROUND",   (0,1), (-1,1), MID_BG),
        ("TEXTCOLOR",    (0,0), (-1,0), MUTED),
        ("FONTNAME",     (0,0), (-1,0), "Helvetica"),
        ("FONTSIZE",     (0,0), (-1,0), 7.5),
        ("ALIGN",        (0,0), (-1,-1), "CENTER"),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
        ("ROWPADDING",   (0,0), (-1,-1), 10),
        ("LINEBELOW",    (0,0), (-1,0), 0.5, BRAND),
        ("GRID",         (0,0), (-1,-1), 0.3, LIGHT_BG),
    ]))
    story.append(kpi_t)
    story.append(Spacer(1, 12))

    # ── SECTION 1: Company Overview ───────────────────────────────────────────
    story.append(Paragraph("Company Overview", H1))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BRAND))
    story.append(Spacer(1, 6))

    ov_data = [
        ["Market Cap",       _fmt(fund.get("mcap") or fund.get("market_cap"), " Cr"),
         "P/E Ratio",        _fmt(fund.get("pe"))],
        ["52W High",         _money(price.get("week52_high")),
         "52W Low",          _money(price.get("week52_low"))],
        ["Day High",         _money(price.get("high")),
         "Day Low",          _money(price.get("low"))],
        ["Dividend Yield",   _fmt(fund.get("dividend"), "%"),
         "Beta",             _fmt(fund.get("beta"))],
        ["Volume",           f"{int(price.get('volume',0)):,}" if price.get('volume') else "—",
         "Exchange",         exch],
    ]
    ov_t = Table(ov_data, colWidths=["22%","28%","22%","28%"])
    ov_t.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,-1), MID_BG),
        ("ROWBACKGROUND",(0,0), (-1,-1), [MID_BG, LIGHT_BG]),
        ("TEXTCOLOR",    (0,0), (-1,-1), TEXT),
        ("FONTNAME",     (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTNAME",     (2,0), (2,-1), "Helvetica-Bold"),
        ("TEXTCOLOR",    (0,0), (0,-1), MUTED),
        ("TEXTCOLOR",    (2,0), (2,-1), MUTED),
        ("FONTSIZE",     (0,0), (-1,-1), 9),
        ("ROWPADDING",   (0,0), (-1,-1), 7),
        ("GRID",         (0,0), (-1,-1), 0.3, LIGHT_BG),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
    ]))
    story.append(ov_t)
    story.append(Spacer(1, 12))

    # ── SECTION 2: Fundamentals ───────────────────────────────────────────────
    story.append(Paragraph("Fundamental Analysis", H1))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BRAND))
    story.append(Spacer(1, 6))

    fd = [
        ["Metric",               "Value",       "Metric",          "Value"],
        ["P/E Ratio",            _fmt(fund.get("pe")),     "P/B Ratio",   _fmt(fund.get("pb"))],
        ["EPS (TTM)",            _fmt(fund.get("eps")),    "ROE",         _pct(fund.get("roe"))],
        ["ROCE",                 _pct(fund.get("roce")),   "Debt/Equity", _fmt(fund.get("debt_equity"))],
        ["Profit Margin",        _pct(fund.get("profit_margin")), "Current Ratio", _fmt(fund.get("current_ratio"))],
        ["Revenue Growth (YoY)", _pct(fund.get("rev_growth") or fund.get("revenue_growth")),
         "Promoter Stake",       _pct(fund.get("promoter"))],
        ["FII Holding",          _pct(fund.get("fii")),    "DII Holding", _pct(fund.get("dii"))],
    ]
    ft = Table(fd, colWidths=["28%","22%","28%","22%"])
    ft.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,0), LIGHT_BG),
        ("BACKGROUND",   (0,1), (-1,-1), MID_BG),
        ("TEXTCOLOR",    (0,0), (-1,0), MUTED),
        ("TEXTCOLOR",    (0,1), (0,-1), MUTED),
        ("TEXTCOLOR",    (2,1), (2,-1), MUTED),
        ("TEXTCOLOR",    (1,1), (1,-1), WHITE),
        ("TEXTCOLOR",    (3,1), (3,-1), WHITE),
        ("FONTNAME",     (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",     (0,0), (-1,-1), 9),
        ("ROWPADDING",   (0,0), (-1,-1), 7),
        ("GRID",         (0,0), (-1,-1), 0.3, LIGHT_BG),
        ("ALIGN",        (1,0), (1,-1), "RIGHT"),
        ("ALIGN",        (3,0), (3,-1), "RIGHT"),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
    ]))
    story.append(ft)
    story.append(Spacer(1, 12))

    # ── SECTION 3: Technical Analysis ─────────────────────────────────────────
    story.append(Paragraph("Technical Analysis", H1))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BRAND))
    story.append(Spacer(1, 6))

    td = [
        ["Indicator",  "Value",              "Indicator",   "Value"],
        ["RSI (14)",   _fmt(tech.get("rsi")), "MACD Hist",  _fmt(tech.get("macd_hist"))],
        ["ADX",        _fmt(tech.get("adx")), "ATR",        _money(tech.get("atr"))],
        ["MA 20",      _money(tech.get("ma20")), "MA 50",   _money(tech.get("ma50"))],
        ["MA 200",     _money(tech.get("ma200")), "BB Upper", _money(tech.get("bb_upper"))],
        ["BB Lower",   _money(tech.get("bb_lower")), "SAR",  _money(tech.get("sar"))],
        ["Support",    _money(tech.get("support")), "Resistance", _money(tech.get("resistance"))],
        ["Stop Loss",  _money(tech.get("stop_loss")), "Target 1", _money(tech.get("target1"))],
        ["Target 2",   _money(tech.get("target2")), "Bull Score", _fmt(tech.get("bull_score"),"%")],
        ["Golden Cross", "Yes" if tech.get("golden_cross") else "No",
         "Stochastic %K", _fmt(tech.get("stoch_k"))],
    ]
    tt = Table(td, colWidths=["28%","22%","28%","22%"])
    tt.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,0), LIGHT_BG),
        ("BACKGROUND",   (0,1), (-1,-1), MID_BG),
        ("TEXTCOLOR",    (0,0), (-1,0), MUTED),
        ("TEXTCOLOR",    (0,1), (0,-1), MUTED),
        ("TEXTCOLOR",    (2,1), (2,-1), MUTED),
        ("TEXTCOLOR",    (1,1), (1,-1), WHITE),
        ("TEXTCOLOR",    (3,1), (3,-1), WHITE),
        ("FONTNAME",     (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE",     (0,0), (-1,-1), 9),
        ("ROWPADDING",   (0,0), (-1,-1), 7),
        ("GRID",         (0,0), (-1,-1), 0.3, LIGHT_BG),
        ("ALIGN",        (1,0), (1,-1), "RIGHT"),
        ("ALIGN",        (3,0), (3,-1), "RIGHT"),
        ("VALIGN",       (0,0), (-1,-1), "MIDDLE"),
    ]))
    story.append(tt)
    story.append(Spacer(1, 12))

    # Patterns
    patterns = tech.get("patterns", [])
    if patterns:
        story.append(Paragraph("Candlestick Patterns Detected:", H2))
        story.append(Paragraph(", ".join(patterns), NORMAL))
        story.append(Spacer(1, 8))

    # ── SECTION 4: AI Verdict ─────────────────────────────────────────────────
    story.append(Paragraph("AI Verdict & Price Targets", H1))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BRAND))
    story.append(Spacer(1, 6))

    v_signal = verd.get("verdict") or verd.get("signal") or sig
    v_conf   = verd.get("confidence", "N/A")
    v_risk   = verd.get("risk", "N/A")
    v_for    = verd.get("best_for", "N/A")

    vd_data = [
        [Paragraph(f"<b>{v_signal}</b>",
                   S("VS", fontSize=20, textColor=_sig_color(v_signal),
                     fontName="Helvetica-Bold", alignment=TA_CENTER)),
         Paragraph(
             f"<b>Confidence:</b> {v_conf}<br/>"
             f"<b>Risk:</b> {v_risk}<br/>"
             f"<b>Best for:</b> {v_for}<br/>"
             f"<b>3M Target:</b> {_money(verd.get('target_3m'))}<br/>"
             f"<b>12M Target:</b> {_money(verd.get('target_12m'))}<br/>"
             f"<b>Stop Loss:</b> {_money(verd.get('stop_loss') or tech.get('stop_loss'))}",
             S("VD", fontSize=9, textColor=TEXT, fontName="Helvetica", leading=14))
        ]
    ]
    vdt = Table(vd_data, colWidths=["25%","75%"])
    vdt.setStyle(TableStyle([
        ("BACKGROUND",  (0,0), (-1,-1), MID_BG),
        ("VALIGN",      (0,0), (-1,-1), "MIDDLE"),
        ("ROWPADDING",  (0,0), (-1,-1), 12),
        ("LINEAFTER",   (0,0), (0,-1), 1, BRAND),
        ("GRID",        (0,0), (-1,-1), 0.3, LIGHT_BG),
    ]))
    story.append(vdt)
    story.append(Spacer(1, 10))

    # AI Reasoning
    reasoning = verd.get("reasoning", [])
    full_text  = verd.get("full_text", "")
    if reasoning:
        story.append(Paragraph("AI Reasoning:", H2))
        for line in reasoning:
            story.append(Paragraph(f"• {line}", NORMAL))
        story.append(Spacer(1, 6))
    elif full_text:
        story.append(Paragraph("AI Analysis:", H2))
        for para in full_text.split("\n"):
            if para.strip():
                story.append(Paragraph(para.strip(), NORMAL))
        story.append(Spacer(1, 6))

    # ── SECTION 5: Latest News ────────────────────────────────────────────────
    if news:
        story.append(PageBreak())
        story.append(Paragraph("Latest News & Sentiment", H1))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BRAND))
        story.append(Spacer(1, 6))
        for item in (news[:8] if isinstance(news, list) else []):
            if isinstance(item, dict):
                title    = item.get("title", "")
                source   = item.get("source", "")
                pub_date = item.get("published", "") or item.get("date", "")
                if title:
                    story.append(Paragraph(
                        f"<b>{title}</b><br/>"
                        f"<font size='8' color='#5a7090'>{source}  {pub_date}</font>",
                        S("NI", fontSize=9, textColor=TEXT, fontName="Helvetica", leading=14, spaceAfter=6)
                    ))
                    story.append(HRFlowable(width="100%", thickness=0.3, color=LIGHT_BG))

    # ── FOOTER DISCLAIMER ─────────────────────────────────────────────────────
    story.append(Spacer(1, 20))
    story.append(HRFlowable(width="100%", thickness=0.5, color=MUTED))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        "DISCLAIMER: This report is generated by Kevin Kataria Stock Intelligence Pro for informational purposes only. "
        "It does not constitute investment advice. Past performance is not indicative of future results. "
        "Always conduct your own research and consult a SEBI-registered financial advisor before investing. "
        f"Generated: {ts}",
        DISCLAIMER
    ))

    def _page_footer(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(MUTED)
        canvas.setFont("Helvetica", 7)
        canvas.drawString(1.8*cm, 1.2*cm, f"Kevin Kataria Stock Intelligence Pro — {sym}")
        canvas.drawRightString(A4[0] - 1.8*cm, 1.2*cm, f"Page {doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=_page_footer, onLaterPages=_page_footer)
    buf.seek(0)
    return buf


# ═════════════════════════════════════════════════════════════════════════════
# DOCX GENERATOR
# ═════════════════════════════════════════════════════════════════════════════

def generate_docx(context: dict) -> io.BytesIO:
    """Build a Word .docx report from the stock context dict."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    sym  = context.get("symbol",   "N/A")
    name = context.get("name",     sym)
    exch = context.get("exchange", "NSE")
    ts   = datetime.datetime.now().strftime("%d %b %Y, %I:%M %p")
    price= context.get("price", {}) or {}
    tech = context.get("technical", {}) or {}
    fund = context.get("fundamental", {}) or {}
    verd = context.get("verdict", {}) or {}
    news = context.get("news", [])

    def hex_rgb(hex_color):
        h = hex_color.lstrip("#")
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    BRAND_RGB = hex_rgb("38bdf8")
    TEXT_RGB  = hex_rgb("1a2840")
    MUTED_RGB = hex_rgb("5a7090")

    def heading(text, level=1):
        p = doc.add_heading(text, level)
        p.runs[0].font.color.rgb = BRAND_RGB
        return p

    def kv(table, r, c1, v1, c2="", v2=""):
        cells = table.rows[r].cells
        cells[0].text = c1
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[0].paragraphs[0].runs[0].font.color.rgb = MUTED_RGB
        cells[1].text = v1
        if c2:
            cells[2].text = c2
            cells[2].paragraphs[0].runs[0].font.bold = True
            cells[2].paragraphs[0].runs[0].font.color.rgb = MUTED_RGB
        if v2:
            cells[3].text = v2

    # ── Title ─────────────────────────────────────────────────────────────────
    title = doc.add_heading(f"{name} ({sym})", 0)
    title.runs[0].font.color.rgb = BRAND_RGB
    p = doc.add_paragraph(f"{exch} · Report generated {ts}")
    p.runs[0].font.color.rgb = MUTED_RGB
    p.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── Price Summary ─────────────────────────────────────────────────────────
    heading("Price & Signal Summary")
    t = doc.add_table(rows=2, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(["Price", "Day Change", "Signal", "Bull Score"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = MUTED_RGB
    row = t.rows[1].cells
    row[0].text = _money(price.get("price"))
    row[1].text = _pct(price.get("pct") or price.get("change_pct"))
    row[2].text = tech.get("signal", verd.get("verdict","HOLD"))
    row[3].text = _fmt(tech.get("bull_score"), "%")
    doc.add_paragraph()

    # ── Company Overview ──────────────────────────────────────────────────────
    heading("Company Overview")
    t2 = doc.add_table(rows=5, cols=4)
    t2.style = "Table Grid"
    data = [
        ("Market Cap", _fmt(fund.get("mcap"),"Cr"), "P/E Ratio", _fmt(fund.get("pe"))),
        ("52W High",   _money(price.get("week52_high")), "52W Low", _money(price.get("week52_low"))),
        ("Day High",   _money(price.get("high")),  "Day Low",     _money(price.get("low"))),
        ("Div Yield",  _pct(fund.get("dividend")), "Beta",         _fmt(fund.get("beta"))),
        ("Volume",     f"{int(price.get('volume',0)):,}" if price.get("volume") else "—",
         "Exchange",   exch),
    ]
    for i, row_data in enumerate(data):
        kv(t2, i, *row_data)
    doc.add_paragraph()

    # ── Fundamentals ──────────────────────────────────────────────────────────
    heading("Fundamental Analysis")
    t3 = doc.add_table(rows=7, cols=4)
    t3.style = "Table Grid"
    fd = [
        ("P/E Ratio",        _fmt(fund.get("pe")),   "P/B Ratio",    _fmt(fund.get("pb"))),
        ("EPS (TTM)",        _fmt(fund.get("eps")),  "ROE",           _pct(fund.get("roe"))),
        ("ROCE",             _pct(fund.get("roce")), "Debt/Equity",   _fmt(fund.get("debt_equity"))),
        ("Profit Margin",    _pct(fund.get("profit_margin")), "Current Ratio", _fmt(fund.get("current_ratio"))),
        ("Revenue Growth",   _pct(fund.get("rev_growth")), "Promoter %",  _pct(fund.get("promoter"))),
        ("FII Holding",      _pct(fund.get("fii")),  "DII Holding",  _pct(fund.get("dii"))),
        ("Source",           str(fund.get("source","Yahoo Finance")), "", ""),
    ]
    for i, row_data in enumerate(fd):
        kv(t3, i, *row_data)
    doc.add_paragraph()

    # ── Technical Analysis ────────────────────────────────────────────────────
    heading("Technical Analysis")
    t4 = doc.add_table(rows=10, cols=4)
    t4.style = "Table Grid"
    td = [
        ("RSI (14)",     _fmt(tech.get("rsi")),          "MACD Hist",  _fmt(tech.get("macd_hist"))),
        ("ADX",          _fmt(tech.get("adx")),           "ATR",        _money(tech.get("atr"))),
        ("MA 20",        _money(tech.get("ma20")),         "MA 50",      _money(tech.get("ma50"))),
        ("MA 200",       _money(tech.get("ma200")),        "BB Upper",   _money(tech.get("bb_upper"))),
        ("BB Lower",     _money(tech.get("bb_lower")),     "SAR",        _money(tech.get("sar"))),
        ("Support",      _money(tech.get("support")),      "Resistance", _money(tech.get("resistance"))),
        ("Stop Loss",    _money(tech.get("stop_loss")),    "Target 1",   _money(tech.get("target1"))),
        ("Target 2",     _money(tech.get("target2")),      "Bull Score", _fmt(tech.get("bull_score"),"%")),
        ("Golden Cross", "Yes" if tech.get("golden_cross") else "No",
         "Stochastic",   _fmt(tech.get("stoch_k"))),
        ("Patterns",     ", ".join(tech.get("patterns",[]) or ["None detected"]), "", ""),
    ]
    for i, row_data in enumerate(td):
        kv(t4, i, *row_data)
    doc.add_paragraph()

    # ── AI Verdict ────────────────────────────────────────────────────────────
    heading("AI Verdict")
    v_sig = verd.get("verdict") or verd.get("signal") or "HOLD"
    p = doc.add_paragraph()
    r = p.add_run(v_sig)
    r.bold = True; r.font.size = Pt(18)
    if "BUY" in v_sig.upper():    r.font.color.rgb = hex_rgb("10b981")
    elif "SELL" in v_sig.upper(): r.font.color.rgb = hex_rgb("f87171")
    else:                          r.font.color.rgb = hex_rgb("fbbf24")

    details = doc.add_paragraph()
    details.add_run(f"Confidence: {verd.get('confidence','N/A')}  |  "
                    f"Risk: {verd.get('risk','N/A')}  |  "
                    f"Best for: {verd.get('best_for','N/A')}\n"
                    f"3M Target: {_money(verd.get('target_3m'))}  |  "
                    f"12M Target: {_money(verd.get('target_12m'))}  |  "
                    f"Stop Loss: {_money(verd.get('stop_loss') or tech.get('stop_loss'))}")

    reasoning = verd.get("reasoning", [])
    if reasoning:
        doc.add_heading("Reasoning", level=2)
        for line in reasoning:
            doc.add_paragraph(f"• {line}", style="List Bullet")
    elif verd.get("full_text"):
        doc.add_heading("AI Analysis", level=2)
        for para in (verd.get("full_text","")).split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

    # ── News ──────────────────────────────────────────────────────────────────
    if news:
        doc.add_page_break()
        heading("Latest News")
        for item in (news[:8] if isinstance(news, list) else []):
            if isinstance(item, dict) and item.get("title"):
                p = doc.add_paragraph()
                r = p.add_run(item["title"])
                r.bold = True
                p.add_run(f"\n{item.get('source','')}  {item.get('published','') or item.get('date','')}")

    # ── Disclaimer ────────────────────────────────────────────────────────────
    doc.add_page_break()
    p = doc.add_paragraph(
        "DISCLAIMER: This report is generated by Kevin Kataria Stock Intelligence Pro for "
        "informational and educational purposes only. It does not constitute investment advice, "
        "solicitation, or recommendation to buy or sell any security. "
        "Always conduct your own research and consult a SEBI-registered investment advisor. "
        f"Generated: {ts}"
    )
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = MUTED_RGB

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═════════════════════════════════════════════════════════════════════════════
# FLASK ROUTE REGISTRATION
# ═════════════════════════════════════════════════════════════════════════════

def register_report_routes(app, build_stock_context_fn):
    """
    Call after app = Flask(...):
        from report_generator import register_report_routes
        register_report_routes(app, build_stock_context)
    """
    @app.route("/api/report/<symbol>")
    def download_report(symbol):
        sym    = symbol.upper().strip()
        fmt    = (request.args.get("format") or "pdf").lower()
        if fmt not in ("pdf", "docx"):
            return jsonify({"error": "format must be 'pdf' or 'docx'"}), 400

        # Build context
        try:
            ctx = build_stock_context_fn(sym)
        except Exception as e:
            ctx = {"symbol": sym, "error": str(e)}

        # Flatten nested data for easy access
        context = {
            "symbol":      sym,
            "name":        ctx.get("name", sym),
            "exchange":    ctx.get("exchange", "NSE"),
            "price":       ctx.get("price", {}),
            "technical":   ctx.get("technical", {}),
            "fundamental": ctx.get("fundamental", {}),
            "verdict":     ctx.get("verdict", {}),
            "news":        ctx.get("news", []),
        }

        safe_sym = re.sub(r"[^A-Z0-9_.-]", "", sym)
        date_str = datetime.datetime.now().strftime("%Y%m%d")

        if fmt == "pdf":
            if not REPORTLAB_OK:
                return jsonify({"error": "reportlab not installed. Add to requirements.txt"}), 500
            try:
                buf      = generate_pdf(context)
                filename = f"{safe_sym}_report_{date_str}.pdf"
                return send_file(buf, mimetype="application/pdf",
                                 as_attachment=True, download_name=filename)
            except Exception as e:
                return jsonify({"error": f"PDF generation failed: {str(e)[:200]}"}), 500

        else:  # docx
            if not DOCX_OK:
                return jsonify({"error": "python-docx not installed. Add to requirements.txt"}), 500
            try:
                buf      = generate_docx(context)
                filename = f"{safe_sym}_report_{date_str}.docx"
                return send_file(buf,
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    as_attachment=True, download_name=filename)
            except Exception as e:
                return jsonify({"error": f"DOCX generation failed: {str(e)[:200]}"}), 500

import re  # needed inside flask route
